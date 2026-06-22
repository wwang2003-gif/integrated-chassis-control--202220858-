"""build_docx_report.py — model_calibration_report.md → .docx 변환

종합 모델링 보고서 (Word 형식) 생성:
- 제목 페이지 + 목차 placeholder
- 헤딩, 단락, 표, 코드/수식 블록, 이미지 처리
- docs/figures/ PNG들 (FBD + 비교 plot)을 자동 임베드
- 참고문헌 정리

Usage:
    python scripts/utils/build_docx_report.py
출력: docs/Model_Calibration_Report.docx
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Resolve project root from script location
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent
DOCS_DIR = PROJECT_ROOT / "docs"
FIG_DIR = DOCS_DIR / "figures"

# Build BOTH reports
REPORTS = [
    {
        "md": DOCS_DIR / "model_calibration_report.md",
        "docx": DOCS_DIR / "Model_Calibration_Report.docx",
        "title": "Integrated Chassis Control",
        "subtitle": "Plant Model Calibration Report",
        "meta_lines": [
            "Reference Vehicle: BMW 5 Series (CarMaker 15)",
            "Plant Models: Bicycle · 3-DOF · 7-DOF · 14-DOF",
            "Tire Models: simple_mf · full_mf (Pacejka MF 5.2) · rt_proxy",
            "Integrator: RK4 (3/7/14-DOF) · Euler (Bicycle LTI)",
            "Report Date: 2026-05-23",
        ],
    },
    {
        "md": DOCS_DIR / "icc_test_protocol.md",
        "docx": DOCS_DIR / "ICC_Test_Protocol.docx",
        "title": "Integrated Chassis Control",
        "subtitle": "Test Protocol — Standard Scenarios · KPIs · Implementation Plan",
        "meta_lines": [
            "Scope: Longitudinal + Lateral + Vertical Integrated Chassis Control",
            "Categories: A (Lateral) · B (Longitudinal) · C (Vertical) · D (Integration)",
            "Total Scenarios: 26 across 4 categories",
            "References: ISO, UN-R, NHTSA FMVSS",
            "Report Date: 2026-05-23",
        ],
    },
]


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "808080")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_inline_runs(paragraph, text):
    """텍스트의 인라인 마크다운(**bold**, *italic*, `code`, [link](url))을 docx run으로 처리."""
    # 이미지 마크다운은 호출자에서 미리 분리되므로 여기 안 옴
    # ![alt](src)
    # 처리 패턴: code(`), bold(**), italic(*), link([text](url))
    pattern = re.compile(
        r"(`[^`]+`)"
        r"|(\*\*[^*]+\*\*)"
        r"|(\*[^*]+\*)"
        r"|(\[[^\]]+\]\([^)]+\))"
    )
    pos = 0
    for m in pattern.finditer(text):
        start, end = m.span()
        if start > pos:
            paragraph.add_run(text[pos:start])
        token = m.group(0)
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("[") and "](" in token:
            mm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            link_text, link_url = mm.group(1), mm.group(2)
            run = paragraph.add_run(link_text)
            run.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
            run.underline = False  # link target is local file path; keep clean
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_markdown_lines(md_text):
    """기본적인 마크다운 라인 단위 파서. 블록 단위로 (type, content) 튜플 yield.
    types: 'heading', 'paragraph', 'code', 'table', 'list_item', 'image', 'hr', 'blank'."""
    lines = md_text.splitlines()
    i = 0
    blocks = []
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            blocks.append(("blank", ""))
            i += 1
            continue
        # Code fence
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            blocks.append(("code", "\n".join(code_lines)))
            i += 1
            continue
        # Headings (# .. ######)
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            blocks.append(("heading", (level, m.group(2).strip())))
            i += 1
            continue
        # Horizontal rule
        if re.match(r"^---+$", line.strip()):
            blocks.append(("hr", ""))
            i += 1
            continue
        # Image (standalone line)
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
        if m:
            blocks.append(("image", (m.group(1), m.group(2))))
            i += 1
            continue
        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?(\s*:?-+:?\s*\|)+", lines[i + 1]):
            rows = [line]
            i += 1
            sep = lines[i]
            i += 1
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append(lines[i])
                i += 1
            blocks.append(("table", rows))
            continue
        # List item
        if re.match(r"^\s*[-*]\s+", line):
            blocks.append(("list_item", re.sub(r"^\s*[-*]\s+", "", line)))
            i += 1
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            blocks.append(("list_item", re.sub(r"^\s*\d+\.\s+", "", line)))
            i += 1
            continue
        # Paragraph (merge consecutive non-special lines)
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip():
                break
            if re.match(r"^#{1,6}\s+", nxt):
                break
            if nxt.strip().startswith("```"):
                break
            if re.match(r"^---+$", nxt.strip()):
                break
            if re.match(r"^!\[", nxt.strip()):
                break
            if re.match(r"^\s*[-*]\s+", nxt):
                break
            if "|" in nxt and i + 1 < len(lines) and re.match(r"^\s*\|?(\s*:?-+:?\s*\|)+", lines[i + 1]):
                break
            para_lines.append(nxt)
            i += 1
        blocks.append(("paragraph", " ".join(p.strip() for p in para_lines)))
    return blocks


def parse_table_rows(rows):
    """마크다운 표 → 2D list."""
    def parse_row(r):
        # split by '|', strip leading/trailing empty cells
        cells = [c.strip() for c in r.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        return cells
    return [parse_row(r) for r in rows]


def resolve_image_src(src):
    """relative path → absolute path. 'figures/foo.png' → DOCS_DIR/figures/foo.png."""
    p = Path(src)
    if p.is_absolute():
        return p
    return (DOCS_DIR / p).resolve()


def build_one_report(md_path, out_path, title_text, subtitle_text, meta_lines):
    md_text = md_path.read_text(encoding="utf-8")
    blocks = parse_markdown_lines(md_text)

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # === Title page ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(26)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(subtitle_text)
    sub_run.bold = True
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("\n")
    for line in meta_lines:
        meta.add_run(line + "\n").italic = True

    doc.add_page_break()

    # === Body render ===
    list_counter = 0
    in_list = False

    for block_type, content in blocks:
        if block_type == "heading":
            level, text = content
            heading_levels = {1: 0, 2: 1, 3: 2, 4: 3, 5: 4, 6: 5}
            level_docx = heading_levels.get(level, 0)
            h = doc.add_heading(text, level=level_docx)
            # color
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
            in_list = False

        elif block_type == "paragraph":
            p = doc.add_paragraph()
            add_inline_runs(p, content)
            in_list = False

        elif block_type == "code":
            p = doc.add_paragraph()
            # Background
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F2F2F2")
            pPr.append(shd)
            run = p.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            in_list = False

        elif block_type == "image":
            alt, src = content
            img_path = resolve_image_src(src)
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                # determine appropriate width
                width_in = 6.0  # default
                if "fbd_" in src.lower():
                    width_in = 5.5
                elif "compare_" in src.lower():
                    width_in = 6.5
                elif "tire_model_compare" in src.lower():
                    width_in = 6.5
                try:
                    run.add_picture(str(img_path), width=Inches(width_in))
                except Exception as e:
                    p.add_run(f"[IMAGE ERROR: {img_path.name}: {e}]")
                # caption
                if alt:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap.add_run(alt)
                    cap_run.italic = True
                    cap_run.font.size = Pt(10)
                    cap_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            else:
                p = doc.add_paragraph(f"[Missing image: {src}]")
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0xC0, 0x40, 0x40)
            in_list = False

        elif block_type == "table":
            rows = parse_table_rows(content)
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Light Grid Accent 1"
            for ri, row_cells in enumerate(rows):
                for ci in range(ncols):
                    cell = table.rows[ri].cells[ci]
                    val = row_cells[ci] if ci < len(row_cells) else ""
                    cell.text = ""
                    para = cell.paragraphs[0]
                    add_inline_runs(para, val)
                    if ri == 0:
                        for rr in para.runs:
                            rr.bold = True
                    set_cell_borders(cell)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            doc.add_paragraph()  # gap after table
            in_list = False

        elif block_type == "list_item":
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, content)
            in_list = True

        elif block_type == "hr":
            hr_p = doc.add_paragraph()
            add_horizontal_line(hr_p)
            in_list = False

        elif block_type == "blank":
            # Skip extra blank to avoid double-spacing
            in_list = False
            continue

    # === Save ===
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    size_kb = out_path.stat().st_size / 1024
    print(f"Saved: {out_path}  ({size_kb:.1f} KB)")


def main():
    for r in REPORTS:
        if not r["md"].exists():
            print(f"[skip] markdown not found: {r['md']}")
            continue
        build_one_report(r["md"], r["docx"], r["title"], r["subtitle"], r["meta_lines"])


if __name__ == "__main__":
    main()
